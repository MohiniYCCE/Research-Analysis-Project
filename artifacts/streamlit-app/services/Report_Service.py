import pandas as pd
import numpy as np
import io
import os
from datetime import datetime
from pathlib import Path

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable
    )
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def generate_pdf_report(df: pd.DataFrame, filename: str, insights: list = None) -> bytes:
    """Generate a PDF analysis report."""
    if not REPORTLAB_AVAILABLE:
        return b""

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=50, leftMargin=50,
                             topMargin=60, bottomMargin=50)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle("Title", parent=styles["Title"],
                                  fontSize=20, textColor=colors.HexColor("#1d4ed8"),
                                  spaceAfter=6)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"],
                               fontSize=13, textColor=colors.HexColor("#1e40af"),
                               spaceBefore=14, spaceAfter=4)
    body_style = styles["BodyText"]
    body_style.fontSize = 9

    story = []

    # Header
    story.append(Paragraph("Statyx AI — Analysis Report", title_style))
    story.append(Paragraph(f"Dataset: {filename}", styles["Normal"]))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["Normal"]))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#dbeafe")))
    story.append(Spacer(1, 12))

    # Dataset Overview
    story.append(Paragraph("Dataset Overview", h2_style))
    overview_data = [
        ["Metric", "Value"],
        ["Total Rows", f"{len(df):,}"],
        ["Total Columns", f"{len(df.columns)}"],
        ["Numeric Columns", f"{len(df.select_dtypes(include=[np.number]).columns)}"],
        ["Categorical Columns", f"{len(df.select_dtypes(include=['object', 'category']).columns)}"],
        ["Missing Values", f"{df.isnull().sum().sum():,}"],
        ["Duplicate Rows", f"{df.duplicated().sum():,}"],
        ["Memory Usage", f"{df.memory_usage(deep=True).sum() / 1024:.1f} KB"],
    ]
    t = Table(overview_data, colWidths=[3 * inch, 2 * inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1d4ed8")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f0f9ff"), colors.white]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#dbeafe")),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(t)
    story.append(Spacer(1, 14))

    # Descriptive Statistics
    numeric_df = df.select_dtypes(include=[np.number])
    if not numeric_df.empty:
        story.append(Paragraph("Descriptive Statistics", h2_style))
        desc = numeric_df.describe().round(4)
        data = [[""] + list(desc.columns[:8])]
        for row_name in desc.index:
            row = [row_name] + [str(v) for v in desc.loc[row_name].values[:8]]
            data.append(row)

        col_w = [1.2 * inch] + [0.8 * inch] * min(8, len(desc.columns))
        t2 = Table(data, colWidths=col_w)
        t2.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1d4ed8")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#dbeafe")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f0f9ff"), colors.white]),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(t2)
        story.append(Spacer(1, 14))

    # Categorical Summary
    cat_df = df.select_dtypes(include=["object", "category"])
    if not cat_df.empty:
        story.append(Paragraph("Categorical Summary", h2_style))
        for col in cat_df.columns:
            counts = df[col].value_counts(dropna=False).head(10)
            total = len(df)
            data = [["Category", "Count", "Percentage"]]
            for value, count in counts.items():
                label = "Missing" if pd.isna(value) else str(value)
                data.append([label, str(count), f"{round(count / total * 100, 1)}%"])

            t_cat = Table(data, colWidths=[2.4 * inch, 1.1 * inch, 1.1 * inch])
            t_cat.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1d4ed8")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#dbeafe")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f0f9ff"), colors.white]),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(t_cat)
            story.append(Spacer(1, 12))
    
    # AI Insights
    if insights:
        story.append(Paragraph("AI-Generated Insights", h2_style))
        for ins in insights:
            story.append(Paragraph(f"• {ins['title']}: {ins['description']}", body_style))
            story.append(Spacer(1, 4))

    # Column Summary
    story.append(PageBreak())
    story.append(Paragraph("Column Summary", h2_style))
    col_data = [["Column", "Type", "Non-Null", "Unique", "Missing %"]]
    for col in df.columns:
        missing_pct = round(df[col].isnull().sum() / len(df) * 100, 1)
        col_data.append([
            col[:30], str(df[col].dtype),
            f"{df[col].count():,}", f"{df[col].nunique():,}", f"{missing_pct}%"
        ])
    t3 = Table(col_data, colWidths=[2.5 * inch, 1 * inch, 1 * inch, 1 * inch, 0.8 * inch])
    t3.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1d4ed8")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#dbeafe")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f0f9ff"), colors.white]),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(t3)

    doc.build(story)
    return buf.getvalue()


def generate_word_report(df: pd.DataFrame, filename: str, insights: list = None) -> bytes:
    """Generate a Word (.docx) analysis report."""
    if not DOCX_AVAILABLE:
        return b""

    doc = Document()

    # Title
    title = doc.add_heading("Statyx AI — Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(29, 78, 216)

    doc.add_paragraph(f"Dataset: {filename}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("─" * 60)

    # Overview
    doc.add_heading("Dataset Overview", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Shading"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Metric"
    hdr_cells[1].text = "Value"

    rows = [
        ("Total Rows", f"{len(df):,}"),
        ("Total Columns", str(len(df.columns))),
        ("Numeric Columns", str(len(df.select_dtypes(include=[np.number]).columns))),
        ("Categorical Columns", str(len(df.select_dtypes(include=["object", "category"]).columns))),
        ("Missing Values", f"{df.isnull().sum().sum():,}"),
        ("Duplicate Rows", f"{df.duplicated().sum():,}"),
    ]
    for k, v in rows:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v

    doc.add_paragraph("")

    # Statistics
    numeric_df = df.select_dtypes(include=[np.number])
    if not numeric_df.empty:
        doc.add_heading("Descriptive Statistics", level=1)
        desc = numeric_df.describe().round(4)
        cols = list(desc.columns[:6])
        t = doc.add_table(rows=1, cols=len(cols) + 1)
        t.style = "Light Shading"
        hdr = t.rows[0].cells
        hdr[0].text = "Statistic"
        for i, c in enumerate(cols):
            hdr[i + 1].text = c
        for idx in desc.index:
            row = t.add_row().cells
            row[0].text = idx
            for i, c in enumerate(cols):
                row[i + 1].text = str(desc.loc[idx, c])
        doc.add_paragraph("")

    # Categorical Summary
    cat_df = df.select_dtypes(include=["object", "category"])
    if not cat_df.empty:
        doc.add_heading("Categorical Summary", level=1)
        for col in cat_df.columns:
            doc.add_heading(col, level=2)
            counts = df[col].value_counts(dropna=False).head(10)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Shading"
            hdr = table.rows[0].cells
            hdr[0].text = "Category"
            hdr[1].text = "Count"
            hdr[2].text = "Percentage"
            total = len(df)
            for value, count in counts.items():
                row = table.add_row().cells
                label = "Missing" if pd.isna(value) else str(value)
                row[0].text = label
                row[1].text = str(count)
                row[2].text = f"{round(count / total * 100, 1)}%"
            doc.add_paragraph("")

    # AI Insights
    if insights:
        doc.add_heading("AI-Generated Insights", level=1)
        for ins in insights:
            p = doc.add_paragraph()
            p.add_run(f"{ins['title']}: ").bold = True
            p.add_run(ins["description"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
